# annotate .docx (inline comments + highlighting)
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import os

def annotate_docx_simple(original_path, issues):
    """
    For each issue related to this document, highlight snippet occurrences and
    append an auto-comment paragraph at end of the document.
    Returns path to reviewed docx.
    """
    doc = Document(original_path)
    lowered_paras = [p.text.lower() for p in doc.paragraphs]
    for issue in issues:
        # only annotate issues that target this doc
        if issue.get("document") and os.path.abspath(issue["document"]) != os.path.abspath(original_path):
            continue
        snippet = (issue.get("snippet") or "").strip()
        labelled = False
        if snippet:
            snippet_lower = snippet.lower()
            for p in doc.paragraphs:
                if snippet_lower in p.text.lower():
                    for run in p.runs:
                        if snippet_lower in run.text.lower():
                            try:
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                labelled = True
                            except Exception:
                                pass
                    if labelled:
                        break
        # append a comment paragraph at the end (simple and visible)
        comment_text = f"AUTO-COMMENT: {issue.get('issue')} Suggestion: {issue.get('suggestion')}"
        doc.add_paragraph(comment_text)

        # optionally, if rag_contexts present, append them too
        if "rag_contexts" in issue and issue["rag_contexts"]:
            doc.add_paragraph("RAG Contexts / citations:")
            for c in issue["rag_contexts"]:
                title = c.get("meta", {}).get("title") or c.get("meta", {}).get("source") or "source"
                excerpt = c.get("text")[:300].replace("\n", " ")
                doc.add_paragraph(f"- {title}: {excerpt}...")

    reviewed_path = original_path.replace(".docx", ".reviewed.docx")
    doc.save(reviewed_path)
    return reviewed_path
