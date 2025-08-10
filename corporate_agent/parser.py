# parse .docx into structured text
from docx import Document
import docx2txt
import os

def parse_docx(path):
    """
    Returns dict:
      - path
      - full_text
      - paragraphs (list)
      - headings (list of (style_name, text))
      - tables (list of 2D lists)
      - doc (Document object) -- optional for annotator
    """
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    headings = []
    for p in doc.paragraphs:
        style_name = getattr(p.style, "name", "")
        if "heading" in style_name.lower():
            headings.append((style_name, p.text))
    tables = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            rows.append([c.text for c in r.cells])
        tables.append(rows)
    # docx2txt fallback gives plain text
    try:
        full_text = docx2txt.process(path)
    except Exception:
        # fallback to joining paragraphs
        full_text = "\n".join(paragraphs)
    return {
        "path": path,
        "full_text": full_text,
        "paragraphs": paragraphs,
        "headings": headings,
        "tables": tables,
        "document_obj": doc
    }